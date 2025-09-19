import re
from pathlib import Path
from typing import List, Dict, Any, Iterable, Optional

import pdfplumber
from docx import Document as DocxDocument
import pandas as pd


def normalize_whitespace(text: str) -> str:
    text = text.replace('\u00A0', ' ')
    text = re.sub(r'[\t\r]+', ' ', text)
    text = re.sub(r'\s+\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


def extract_text_from_pdf(file_path: Path) -> str:
    parts: List[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ''
            if txt:
                parts.append(txt)
    return normalize_whitespace('\n\n'.join(parts))


def extract_text_from_docx(file_path: Path) -> str:
    doc = DocxDocument(str(file_path))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    # tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(' | '.join(cells))
    return normalize_whitespace('\n'.join(parts))


def extract_text_from_excel(file_path: Path) -> str:
    xl = pd.ExcelFile(str(file_path))
    parts: List[str] = []
    for sheet in xl.sheet_names:
        try:
            df = xl.parse(sheet)
            # Convert to markdown-like table for readability
            # 确保df是DataFrame对象
            if isinstance(df, pd.DataFrame):
                df = df.fillna('')
                header = ' | '.join(map(str, df.columns))
                rows = [' | '.join(map(lambda v: str(v), row)) for row in df.values]
                content = f'# Sheet: {sheet}\n{header}\n' + '\n'.join(rows)
                parts.append(content)
            else:
                parts.append(f'# Sheet: {sheet}\n(Unexpected data format)')
        except Exception as e:
            # If sheet parsing fails, add a note
            parts.append(f'# Sheet: {sheet}\n(Failed to parse: {str(e)})')
    return normalize_whitespace('\n\n'.join(parts))


def chunk_text(text: str, max_chars: int = 2000, overlap: int = 200) -> List[Dict[str, Any]]:
    if not text:
        return []
    chunks: List[Dict[str, Any]] = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + max_chars, text_len)
        chunk_text_val = text[start:end]
        # avoid cutting mid-sentence if possible
        if end < text_len:
            last_period = chunk_text_val.rfind('\n')
            if last_period > max_chars // 2:
                end = start + last_period
                chunk_text_val = text[start:end]
        chunks.append({
            'text': chunk_text_val.strip(),
        })
        if end == text_len:
            break
        start = max(0, end - overlap)
    return chunks


def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from plain text files"""
    try:
        content = file_path.read_text(encoding='utf-8')
        return normalize_whitespace(content)
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            content = file_path.read_text(encoding='gbk')
            return normalize_whitespace(content)
        except:
            content = file_path.read_text(encoding='latin-1')
            return normalize_whitespace(content)


def parse_file_to_chunks(file_path: Path, ext: str, max_chars: int = 2000, overlap: int = 200) -> List[Dict[str, Any]]:
    ext = ext.lower().lstrip('.')
    if ext == 'pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == 'docx':
        text = extract_text_from_docx(file_path)
    elif ext in ('xlsx', 'xls'):
        text = extract_text_from_excel(file_path)
    elif ext == 'txt':
        text = extract_text_from_txt(file_path)
    else:
        raise ValueError(f'Unsupported extension: {ext}')

    return chunk_text(text, max_chars=max_chars, overlap=overlap)


